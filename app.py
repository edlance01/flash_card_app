from flask import Flask, render_template, request, redirect, url_for
import docx
import os
import tempfile

app = Flask(__name__)

def get_bullet_style(para):
        numPr = para._element.find(
            ".//w:numPr", namespaces=para._element.nsmap
        )

        return numPr


def extract_terms_and_definitions(doc_path):
    """
    Extracts terms and definitions from a Word document where
    all paragraph styles are 'Normal' and bullet points are
    represented by the bullet point character (•).
    """
    doc = docx.Document(doc_path)
    terms_and_definitions = {}
    term = None
    definition = ""

    for i, paragraph in enumerate(doc.paragraphs):
        print(f"\n\nPARAGRAPH:{paragraph.text}\n")
        if not paragraph.text.strip():
            continue
        if get_bullet_style(paragraph) is not None:
            if term:
                terms_and_definitions[term] = definition.strip()
            term = (
                paragraph.text.split(":")[0].strip()
                if ":" in paragraph.text
                else paragraph.text.strip()
            )

            print(f"\nTERM:{term}")
            definition = (
                paragraph.text.split(":", 1)[1].strip() if ":" in paragraph.text else ""
            )

            print(f"\nDEFINITION:{definition}")

            # Check for blank line after definition
            next_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
            if next_para and not next_para.text.strip():
                break
        elif term:
            definition += " " + paragraph.text

    if term:
        terms_and_definitions[term] = definition.strip()

    return terms_and_definitions

@app.route("/", methods=["GET", "POST"])
def index():
    """
    Main route for the flashcard program.
    """
    if request.method == "POST":
        if "file" not in request.files:
            return redirect(request.url)
        file = request.files["file"]
        if file.filename == "":
            return redirect(request.url)
        if file:
            print("\n***STARTING***")
            # Save the file to a temporary location
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                file.save(tmp.name)
                temp_path = tmp.name
            # Process the saved file
            terms_and_definitions = extract_terms_and_definitions(temp_path)
            # Remove the temporary file
            os.remove(temp_path)
            return redirect(url_for("flashcards", terms=terms_and_definitions))
    return render_template("index.html")


@app.route("/flashcards")
def flashcards():
    """
    Route for displaying the flashcards.
    """
    terms_and_definitions = request.args.get("terms")
    if terms_and_definitions:
        terms_and_definitions = eval(terms_and_definitions)
        return render_template("flashcards.html", flashcards=terms_and_definitions)
    else:
        return redirect(url_for("index"))


if __name__ == "__main__":

    #get_bullet_style("static/input_files/EssentialsOfAI_Module1_Vocab.docx")
    app.run(debug=True)
